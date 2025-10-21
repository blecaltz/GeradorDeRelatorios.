import streamlit as st
import pandas as pd
from docx import Document
import locale
from datetime import datetime
import copy
import os
import io
import zipfile
from docx.shared import RGBColor  
from docx.shared import Pt    

def format_seconds_to_hhmm(seconds):
    """Converte segundos em uma string no formato HH:MM."""
    try:
        total_seconds = int(float(seconds))
        if total_seconds < 0:
            return "00:00"
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02d}:{minutes:02d}"
    except (ValueError, TypeError):
        return "N/A"


def substituir_placeholders(elemento, substituicoes, doc_para_estilos=None):
    for paragrafo in elemento.paragraphs:
        texto_completo = ''.join(run.text for run in paragrafo.runs)

        if "{{tk_sla}}" in texto_completo:
            sla_expirado = substituicoes.get("{{tk_sla_status}}", False)
            texto_sla = "SLA fora do prazo" if sla_expirado else "SLA dentro do prazo"

            for i in range(len(paragrafo.runs)):
                paragrafo.runs[i].text = ''

            run = paragrafo.add_run(texto_sla)

            # Mantém o mesmo tamanho de fonte do parágrafo
            try:
                tamanho_fonte = paragrafo.style.font.size or Pt(6)
                run.font.size = tamanho_fonte
            except Exception:
                run.font.size = Pt(6)

            # Formatação condicional
            if sla_expirado:
                run.font.color.rgb = RGBColor(255, 0, 0)  # vermelho
                run.bold = True
            else:
                run.font.color.rgb = None
                run.bold = False

            if "{{tk_sla}}" in substituicoes:
                del substituicoes["{{tk_sla}}"]
            continue

        # Substituição normal de placeholders
        if any(chave in texto_completo for chave in substituicoes):
            for chave, valor in substituicoes.items():
                if chave in texto_completo:
                    texto_completo = texto_completo.replace(str(chave), str(valor))

            for i in range(len(paragrafo.runs)):
                paragrafo.runs[i].text = ''

            if paragrafo.runs:
                paragrafo.runs[0].text = texto_completo
            else:
                paragrafo.add_run(texto_completo)

    for tabela in elemento.tables:
        for row in tabela.rows:
            for cell in row.cells:
                substituir_placeholders(cell, substituicoes, doc_para_estilos)


def gerar_relatorio_em_memoria(dados_relatorio, tickets_do_cliente, template_stream):
    if isinstance(template_stream, str):
        doc_principal = Document(template_stream)
    else:
        doc_principal = Document(io.BytesIO(template_stream.getvalue()))

    mes_referencia_dt = dados_relatorio['MÊS DE REFERENCIA']

    gestor = str(dados_relatorio.get('GESTOR_DO_CONTRATO', ''))
    fiscal_tecnico = dados_relatorio.get('FISCAL_TECNICO')
    valor_fiscal = "" if pd.isna(fiscal_tecnico) else f"Fiscal técnico: {fiscal_tecnico}"
    aos_cuidados_valor = dados_relatorio.get('AOS_CUIDADOS')
    valor_aos_cuidados = "" if pd.isna(aos_cuidados_valor) else f"A/C: {aos_cuidados_valor}"

    substituicoes_gerais = {
        "{{NOME_DO_CLIENTE}}": str(dados_relatorio.get('CLIENTE', '')),
        "{{NUMERO_CONTRATO}}": str(dados_relatorio.get('N° DO CONTRATO NO SERVICENOW', '')),
        "{{MES_REFERENCIA_MAIUSCULO}}": mes_referencia_dt.strftime('%B de %Y').upper(),
        "{{PERIODO_REFERENCIA}}": f"01 de {mes_referencia_dt.strftime('%B')} a 31 de {mes_referencia_dt.strftime('%B de %Y')}",
        "{{DATA_RELATORIO}}": datetime.now().strftime('%d/%m/%Y %H:%M'),
        "{{DISPONIBILIDADE}}": str(dados_relatorio.get('DISPONIBILIDADE', '')),
        "{{GESTOR_DO_CONTRATO}}": gestor,
        "{{FISCAL_TECNICO}}": valor_fiscal,
        "{{AOS_CUIDADOS}}": valor_aos_cuidados
    }

    substituir_placeholders(doc_principal, substituicoes_gerais, doc_principal)

    if doc_principal.tables:
        tabela_modelo = doc_principal.tables[0]

        if not tickets_do_cliente.empty:
            for index, ticket in tickets_do_cliente.iterrows():
                nova_tabela_xml = copy.deepcopy(tabela_modelo._element)
                doc_principal.element.body.append(nova_tabela_xml)
                nova_tabela_docx = doc_principal.tables[-1]

                def get_value(field_name):
                    value = ticket.get(field_name)
                    return "" if pd.isna(value) else str(value)

                ans_expirado_bool = ticket.get("ANS expirado", False)

                substituicoes_ticket = {
                    "{{tk_numero}}": get_value("Número"),
                    "{{tk_abertura}}": get_value("Aberto"),
                    "{{tk_inicio}}": get_value("Aberto"),
                    "{{tk_conclusao}}": get_value("Atualizado em"),
                    "{{tk_solicitante}}": get_value("Aberto por"),
                    "{{tk_analista}}": get_value("Atribuído a"),
                    "{{tk_referencia}}": get_value("Descrição resumida"),
                    "{{tk_prioridade}}": get_value("Prioridade"),
                    "{{tk_descricao}}": get_value("Descrição"),
                    "{{tk_resolucao}}": get_value("Anotações de encerramento"),
                    "{{tk_sla_status}}": ans_expirado_bool,
                    "{{tk_tempo_inicio}}": get_value("Tempo trabalhado"),
                    "{{tk_tempo_resolucao}}": format_seconds_to_hhmm(ticket.get("Duração dos negócios")),
                    "{{tk_tempo_espera}}": format_seconds_to_hhmm(ticket.get("Tempo em espera")),
                }

                for row in nova_tabela_docx.rows:
                    for cell in row.cells:
                        substituir_placeholders(cell, substituicoes_ticket, doc_principal)

                doc_principal.add_paragraph()
        else:
            doc_principal.add_paragraph("Não houve chamados neste período para este cliente.")

        primeira_tabela_elemento = doc_principal.tables[0]._element
        primeira_tabela_elemento.getparent().remove(primeira_tabela_elemento)

    cliente = dados_relatorio.get('CLIENTE', 'Cliente Desconhecido')
    mes_ano = mes_referencia_dt.strftime('%B %Y')
    nome_arquivo = f"[TLTX] {cliente} Relatório Mensal de Operação - {mes_ano}.docx"

    file_stream = io.BytesIO()
    doc_principal.save(file_stream)
    file_stream.seek(0)

    return nome_arquivo, file_stream


# --- STREAMLIT ---
st.set_page_config(layout="centered", page_title="Gerador de Relatórios")

try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except locale.Error:
    st.warning("Aviso: Não foi possível configurar o idioma para Português (pt_BR). Os meses podem aparecer em inglês.")

st.markdown("""
<style>
    header, footer { visibility: hidden; }
    body { background-color: #1E1E1E !important; }
    .main .block-container { background-color: #1E1E1E; }

    /* Botão Primário (Gerar Relatórios) - Branco/Cinza */
    .stButton>button[kind="primary"] {
        background-color: #FAFAFA !important; /* Fundo branco */
        color: #1E1E1E !important; /* Texto escuro */
        border: 1px solid #CCCCCC !important; /* Borda cinza clara */
    }
    .stButton>button[kind="primary"]:hover {
        background-color: #E0E0E0 !important; /* Cinza um pouco mais escuro no hover */
        color: #1E1E1E !important;
        border: 1px solid #BBBBBB !important;
    }

    /* Botão de Download - Cinza claro */
    .stDownloadButton>button {
        background-color: #CCCCCC !important; /* Fundo cinza claro */
        color: #1E1E1E !important; /* Texto escuro */
        border: 1px solid #AAAAAA !important;
    }
    .stDownloadButton>button:hover {
        background-color: #BBBBBB !important;
        color: #000000 !important;
        border: 1px solid #999999 !important;
    }

    /* Barra de Progresso - Branca */
    .stProgress > div > div > div > div { background-color: #FAFAFA; }

    /* Caixa de Informação (st.info) - Cinza */
    div[data-testid="stInfo"] {
        background-color: #2E2E2E !important;
        color: #A0A0A0 !important;
        border: 1px solid #444444 !important;
    }
    div[data-testid="stInfo"] svg { fill: #A0A0A0 !important; }

    /* Botão de Rádio (Seleção de Template) - Texto branco */
    .stRadio [role="radiogroup"] label span {
        color: #FAFAFA !important;
    }
    .stRadio [role="radio"][aria-checked="true"] div:first-child {
        background-color: #CCCCCC !important; /* Bolinha selecionada cinza */
        border-color: #AAAAAA !important;
    }
    .stRadio [role="radio"][aria-checked="false"] div:first-child {
       border-color: #AAAAAA !important; /* Borda cinza */
    }
</style>
""", unsafe_allow_html=True)

st.title("Gerador Automático de Relatórios")

arquivo_excel_controle = st.file_uploader("1. Envie o arquivo de Controle (.xlsx)", type=["xlsx"])

with st.expander("Ver colunas necessárias para o arquivo de Controle"):
    st.markdown("""
    A planilha de **Controle** deve conter as seguintes colunas:
    - `MÊS DE REFERENCIA`
    - `CLIENTE`
    - `N° DO CONTRATO NO SERVICENOW`
    - `DISPONIBILIDADE`
    - `GESTOR_DO_CONTRATO`
    - `AOS_CUIDADOS` (Opcional)
    - `FISCAL_TECNICO` (Opcional)
    """)

arquivos_tickets = st.file_uploader("2. Envie o(s) arquivo(s) de Chamados (.xlsx)", type=["xlsx"], accept_multiple_files=True)

with st.expander("Ver colunas utilizadas do arquivo de Chamados"):
    st.markdown("""
    A planilha de **Chamados** deve conter, no mínimo, as seguintes colunas:
    - `Empresa`
    - `Número`
    - `Aberto`
    - `Descrição resumida`
    - ... e outras colunas que você deseja usar no template.
    """)

st.write("3. Escolha o Template do Word")

tipo_template = st.radio(
    "Como você quer fornecer o template?",
    ('Usar o template padrão (`modelo_relatorio.docx`)', 'Navegar e escolher um arquivo customizado'),
    label_visibility="collapsed"
)

template_final = None

if tipo_template == 'Navegar e escolher um arquivo customizado':
    with st.expander("Ver placeholders necessários para o template customizado"):
        st.markdown("""
        **Campos Gerais:**
        `{{NOME_DO_CLIENTE}}`, `{{NUMERO_CONTRATO}}`, `{{MES_REFERENCIA_MAIUSCULO}}`,
        `{{PERIODO_REFERENCIA}}`, `{{DATA_RELATORIO}}`, `{{DISPONIBILIDADE}}`,
        `{{GESTOR_DO_CONTRATO}}`, `{{FISCAL_TECNICO}}`, `{{AOS_CUIDADOS}}`

        **Campos da Tabela de Tickets:**
        `{{tk_numero}}`, `{{tk_abertura}}`, `{{tk_inicio}}`, `{{tk_conclusao}}`, `{{tk_solicitante}}`,
        `{{tk_analista}}`, `{{tk_referencia}}`, `{{tk_prioridade}}`, `{{tk_sla}}`,
        `{{tk_tipo_atendimento}}`, `{{tk_tempo_inicio}}`, `{{tk_tempo_resolucao}}`,
        `{{tk_tempo_espera}}`, `{{tk_descricao}}`, `{{tk_resolucao}}`
        """)
    template_final = st.file_uploader("Selecione o arquivo de template (.docx)", type=["docx"])
else:
    template_final = "modelo_relatorio.docx"

if st.button("Gerar e Compactar Relatórios", type="primary", use_container_width=True):
    if not arquivo_excel_controle or not template_final or not arquivos_tickets:
        st.error("Por favor, envie todos os arquivos necessários.")
    else:
        try:
            if isinstance(template_final, str) and not os.path.exists(template_final):
                st.error(f"ERRO: O template padrão '{template_final}' não foi encontrado na pasta do programa.")
            else:
                df_controle = pd.read_excel(arquivo_excel_controle)
                lista_de_dfs_tickets = [pd.read_excel(f) for f in arquivos_tickets]
                df_todos_tickets = pd.concat(lista_de_dfs_tickets, ignore_index=True)

                st.success(f"{len(df_todos_tickets)} chamados carregados de {len(arquivos_tickets)} arquivo(s).")

                zip_buffer = io.BytesIO()
                progress_bar = st.progress(0, text="Iniciando...")

                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for i, (index, linha_controle) in enumerate(df_controle.iterrows()):
                        nome_cliente = linha_controle['CLIENTE']
                        status_text = f"Gerando: {nome_cliente}... ({i+1}/{len(df_controle)})"
                        progress_bar.progress((i + 1) / len(df_controle), text=status_text)

                        tickets_do_cliente = df_todos_tickets[df_todos_tickets['Empresa'] == nome_cliente]

                        nome_arquivo, stream_do_arquivo = gerar_relatorio_em_memoria(
                            linha_controle, tickets_do_cliente, template_final
                        )
                        zf.writestr(nome_arquivo, stream_do_arquivo.read())

                progress_bar.empty()
                st.subheader("✅ Compactação concluída!")
                st.download_button(
                    label="Clique aqui para baixar o arquivo .ZIP",
                    data=zip_buffer.getvalue(),
                    file_name="Relatorios_Gerados.zip",
                    mime="application/zip",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"Ocorreu um erro: {e}")
