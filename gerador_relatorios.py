# --- CÓDIGO CORRIGIDO (VERSÃO 16) COM DEBUG ---

import pandas as pd
from docx import Document
import locale
from datetime import datetime
import copy

DADOS_FALSOS_SERVICE_NOW = {
    "727": [
        {
            "number": "INC0074576", "opened_at": "22/07/2025 13:44:06", "u_start_of_service": "22/07/2025 13:45:56",
            "closed_at": "24/07/2025 12:01:20", "caller_id": "Caixa CETAD 02", "assigned_to": "Matheus Fernandes",
            "short_description": "Alerta Zabbix – Erros de Entrada na Interface", "priority": "4 - Baixa", "sla": "Dentro do prazo",
            "tipo_atendimento": "Presencial", "tempo_inicio_atendimento": "0:01:50", "tempo_resolucao": "02:00:00", "tempo_espera": "00:00:00",
            "description": "Dados do ticket INC0074576 para o contrato 727...",
            "close_notes": "Resolução do ticket INC0074576..."
        }
    ],
    "656": [ 
        {
            "number": "INC0012345", "opened_at": "15/08/2025 10:00:00", "u_start_of_service": "15/08/2025 10:05:00",
            "closed_at": "15/08/2025 18:30:00", "caller_id": "Blue Services", "assigned_to": "Maria Silva",
            "short_description": "Problema de conexão com servidor", "priority": "3 - Média", "sla": "Dentro do prazo",
            "tipo_atendimento": "Remoto", "tempo_inicio_atendimento": "0:05:00", "tempo_resolucao": "08:30:00", "tempo_espera": "00:15:00",
            "description": "Dados do ticket INC0012345 para o contrato 656...",
            "close_notes": "Resolução do ticket INC0012345..."
        }
    ]
}

def buscar_tickets_de_exemplo(numero_contrato):
    print(f"Buscando tickets para o contrato '{numero_contrato}' (tipo: {type(numero_contrato)})...")
    resultado = DADOS_FALSOS_SERVICE_NOW.get(str(numero_contrato), [])
    print(f" -> Encontrados {len(resultado)} tickets.")
    return resultado

def substituir_placeholders(elemento, substituicoes):
    for paragrafo in elemento.paragraphs:
        texto_completo = ''.join(run.text for run in paragrafo.runs)
        if any(chave in texto_completo for chave in substituicoes):
            for chave, valor in substituicoes.items():
                if chave in texto_completo:
                    texto_completo = texto_completo.replace(str(chave), str(valor))
            
            for i in range(len(paragrafo.runs)):
                paragrafo.runs[i].text = ''
            if paragrafo.runs:
                paragrafo.runs[0].text = texto_completo
            else:
                paragrafo.add_run(texto_completo)
    for tabela in elemento.tables:
        for row in tabela.rows:
            for cell in row.cells:
                substituir_placeholders(cell, substituicoes)

def gerar_relatorio_completo(dados_relatorio, tickets, caminho_modelo):
    doc = Document(caminho_modelo)
    
    mes_referencia_dt = dados_relatorio['MÊS DE REFERENCIA']
    substituicoes_gerais = {
        "{{NOME_DO_CLIENTE}}": str(dados_relatorio['CLIENTE']), "{{NUMERO_CONTRATO}}": str(dados_relatorio['N° DO CONTRATO NO SERVICENOW']),
        "{{MES_REFERENCIA_MAIUSCULO}}": mes_referencia_dt.strftime('%B de %Y').upper(),
        "{{PERIODO_REFERENCIA}}": f"01 de {mes_referencia_dt.strftime('%B')} a 31 de {mes_referencia_dt.strftime('%B de %Y')}",
        "{{DATA_RELATORIO}}": datetime.now().strftime('%d/%m/%Y %H:%M'), "{{DISPONIBILIDADE}}": str(dados_relatorio['DISPONIBILIDADE'])
    }

    substituir_placeholders(doc, substituicoes_gerais)

    if doc.tables:
        tabela_modelo = doc.tables[0]
        
        # SÓ FAZ ALGO SE HOUVER TICKETS
        if tickets:
            for ticket in tickets:
                nova_tabela_xml = copy.deepcopy(tabela_modelo._element)
                doc.element.body.append(nova_tabela_xml)
                nova_tabela_docx = doc.tables[-1]

                substituicoes_ticket = {
                    "{{tk_numero}}": ticket.get("number", ""), "{{tk_abertura}}": ticket.get("opened_at", ""),
                    "{{tk_inicio}}": ticket.get("u_start_of_service", ""), "{{tk_conclusao}}": ticket.get("closed_at", ""),
                    "{{tk_solicitante}}": ticket.get("caller_id", ""), "{{tk_analista}}": ticket.get("assigned_to", ""),
                    "{{tk_referencia}}": ticket.get("short_description", ""), "{{tk_prioridade}}": ticket.get("priority", ""),
                    "{{tk_sla}}": ticket.get("sla", ""), "{{tk_tipo_atendimento}}": ticket.get("tipo_atendimento", ""),
                    "{{tk_tempo_inicio}}": ticket.get("tempo_inicio_atendimento", ""), "{{tk_tempo_resolucao}}": ticket.get("tempo_resolucao", ""),
                    "{{tk_tempo_espera}}": ticket.get("tempo_espera", ""),
                    "{{tk_descricao}}": ticket.get("description", ""), "{{tk_resolucao}}": ticket.get("close_notes", "")
                }
                
                for row in nova_tabela_docx.rows:
                    for cell in row.cells:
                        substituir_placeholders(cell, substituicoes_ticket)
                
                doc.add_paragraph()
            
            # CORREÇÃO: A tabela modelo só é removida se tivermos adicionado novas tabelas.
            primeira_tabela_elemento = doc.tables[0]._element
            primeira_tabela_elemento.getparent().remove(primeira_tabela_elemento)
        else:
            # Se não houver tickets, simplesmente remove a tabela modelo para não deixar ela em branco no relatório.
            print("Nenhum ticket encontrado, removendo a tabela modelo do relatório final.")
            primeira_tabela_elemento = doc.tables[0]._element
            primeira_tabela_elemento.getparent().remove(primeira_tabela_elemento)
    else:
        print("Aviso: Nenhuma tabela modelo foi encontrada no documento.")

    nome_arquivo = f"Relatório Final - {dados_relatorio['CLIENTE']} - {mes_referencia_dt.strftime('%B %Y')}.docx"
    doc.save(nome_arquivo)
    print(f"-> Relatório gerado com sucesso: {nome_arquivo}\n")


if __name__ == "__main__":
    try:
        locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
    except locale.Error:
        print("Aviso: Locale 'pt_BR.UTF-8' não encontrado.")

    arquivo_excel = "Controle de relatórios.xlsx"
    arquivo_modelo = "modelo_relatorio.docx"

    try:
        df_controle = pd.read_excel(arquivo_excel)
        print("Arquivo Excel lido com sucesso!")

        for index, linha in df_controle.iterrows():
            print(f"--- Processando linha {index+1} do Excel: {linha['CLIENTE']} ---")
            num_contrato = linha['N° DO CONTRATO NO SERVICENOW']
            tickets = buscar_tickets_de_exemplo(num_contrato)
            
            gerar_relatorio_completo(linha, tickets, arquivo_modelo)
                
    except FileNotFoundError:
        print(f"ERRO: Verifique se os arquivos '{arquivo_excel}' e '{arquivo_modelo}' estão na pasta.")
    except KeyError as e:
        print(f"ERRO: A coluna {e} não foi encontrada no seu arquivo Excel.")
    except Exception as e:
        print(f"Ocorreu um erro inesperado: {e}")